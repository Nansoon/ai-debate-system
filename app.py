import streamlit as st
import openai
import requests
import json
from docx import Document
from modules.utils import save_state, load_state, list_saved_debates, load_config
from modules.critic import get_critique_and_score
import os
import matplotlib.pyplot as plt
import matplotlib as mpl

# 密码保护函数：云端启用，开发本地自动跳过
def check_password():
    if os.getenv("IS_LOCAL", "false").lower() == "true":
        return

    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    if not st.session_state.authenticated:
        st.title("🔐 登录验证")
        pwd = st.text_input("请输入访问密码", type="password")
        if st.button("登录"):
            if pwd == st.secrets["login_password"]:
                st.session_state.authenticated = True
                st.rerun()
            else:
                st.error("密码错误，请重试。")
        st.stop()

# 设置苹果风格主题
st.set_page_config(
    page_title="AI 辩论系统",
    layout="wide",
    page_icon="🤖"
)

check_password()

# 创建必要的目录
os.makedirs("data", exist_ok=True)
os.makedirs("output", exist_ok=True)

# 应用苹果设计风格的CSS - 修复选项框高度和卡片布局问题
st.markdown("""
<style>
    html, body, [class*="css"] {
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", "Helvetica Neue", sans-serif;
        color: #1d1d1f;
        background-color: #f5f5f7;
        line-height: 1.5;
    }

    .title-container {
        background: linear-gradient(135deg, #2c6fbb, #5a4fcf);
        border-radius: 14px;
        padding: 20px 24px;
        margin-bottom: 20px;
        box-shadow: 0 4px 16px rgba(0,0,0,0.1);
        color: white;
    }

    h1 {
        font-weight: 700;
        font-size: 2.2rem;
        letter-spacing: -0.5px;
        margin: 0 0 5px 0;
    }

    .title-container p {
        font-size: 1rem;
        opacity: 0.9;
        margin: 0;
        color: rgba(255, 255, 255, 0.85);
    }

    .card {
        background-color: white;
        border-radius: 12px;
        padding: 16px;
        margin-bottom: 16px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.04);
        border: 1px solid rgba(0,0,0,0.03);
    }

    .card-header {
        display: flex;
        align-items: center;
        gap: 8px;
        margin-bottom: 12px;
    }

    .card-header .icon {
        font-size: 1.4rem;
        display: flex;
        align-items: center;
        justify-content: center;
    }

    .card-header h3 {
        margin: 0;
        font-size: 1.3rem;
        font-weight: 600;
    }

    .stButton>button {
        border-radius: 10px !important;
        padding: 10px 20px !important;
        font-size: 1rem !important;
        font-weight: 500 !important;
        transition: all 0.2s ease;
        border: none;
        box-shadow: 0 2px 6px rgba(0,113,227,0.1);
    }

    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,113,227,0.2);
    }

    .primary-button {
        background: linear-gradient(135deg, #2c6fbb, #5a4fcf);
        color: white;
    }

    .secondary-button {
        background: rgba(0,0,0,0.03);
        color: #1d1d1f;
        border: 1px solid #d2d2d7;
    }

    .stTextInput>div>div>input,
    .stSelectbox>div>div>div {
        border-radius: 10px !important;
        padding: 10px 14px !important;
        font-size: 1rem !important;
        line-height: 1.4 !important;
        height: auto !important;
    }

    .stCheckbox label,
    .stRadio label {
        min-height: 32px;
        display: flex;
        align-items: center;
        padding-top: 6px;
        padding-bottom: 6px;
    }

    .stMarkdown h3 {
        font-size: 1.2rem;
        font-weight: 600;
        color: #1d1d1f;
        margin-top: 0;
    }

    .round-container {
        background-color: white;
        border-radius: 12px;
        padding: 16px;
        margin: 16px 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.04);
        border-left: 4px solid #2c6fbb;
    }

    .round-header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        padding-bottom: 8px;
        margin-bottom: 12px;
        border-bottom: 1px solid #e0e0e0;
    }

    .round-title {
        font-size: 1.2rem;
        font-weight: 600;
    }

    .round-subtitle {
        font-size: 0.95rem;
        color: #86868b;
    }

    .debate-content h4 {
        font-size: 1.1rem;
        margin-bottom: 8px;
        color: #2c6fbb;
    }

    .score-container {
        background: linear-gradient(135deg, #f5f5f7, #ffffff);
        border-radius: 12px;
        padding: 16px;
        margin: 20px 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.03);
        border: 1px solid rgba(0,0,0,0.02);
    }

    .score-header {
        font-size: 1.4rem;
        font-weight: 600;
        margin-bottom: 16px;
    }

    .score-card {
        background-color: white;
        border-radius: 10px;
        padding: 16px;
        text-align: center;
        box-shadow: 0 1px 4px rgba(0,0,0,0.03);
    }

    .score-value {
        font-size: 2.4rem;
        font-weight: 700;
        margin: 10px 0;
    }

    .chatgpt-score .score-value { color: #2c6fbb; }
    .deepseek-score .score-value { color: #5a4fcf; }

    .score-label {
        font-size: 1rem;
        color: #86868b;
    }

    .divider {
        height: 1px;
        background: #e0e0e0;
        margin: 24px 0;
    }

    .empty-state {
        text-align: center;
        padding: 30px 20px;
        color: #86868b;
        font-size: 1rem;
    }

    .empty-state-icon {
        font-size: 2.5rem;
        margin-bottom: 10px;
        opacity: 0.3;
    }

    .debate-button-container {
        display: flex;
        justify-content: center;
        margin: 30px 0;
    }

    .style-card, .debate-card {
        height: auto !important;
        overflow-y: auto !important;
    }

    .footer {
        text-align: center;
        padding: 16px 0;
        color: #86868b;
        font-size: 0.9rem;
    }

    @media (max-width: 768px) {
        .title-container {
            padding: 16px;
        }

        h1 {
            font-size: 1.8rem;
        }

        .card {
            padding: 14px;
        }

        .stButton>button {
            padding: 10px 16px !important;
        }
    }
</style>

""", unsafe_allow_html=True)

# 设置Matplotlib为苹果风格
mpl.rcParams['figure.facecolor'] = 'white'
mpl.rcParams['axes.facecolor'] = 'white'
mpl.rcParams['axes.edgecolor'] = '#d2d2d7'
mpl.rcParams['grid.color'] = '#f0f0f0'
mpl.rcParams['text.color'] = '#1d1d1f'
mpl.rcParams['xtick.color'] = '#86868b'
mpl.rcParams['ytick.color'] = '#86868b'
mpl.rcParams['font.family'] = 'sans-serif'

# 标题区域 - 优化版 (修复颜色对比问题)
st.markdown(
    """
    <div class="title-container">
        <h1>AI 辩论系统</h1>
        <p>ChatGPT 与 DeepSeek AI 的智慧辩论平台</p>
    </div>
    """, 
    unsafe_allow_html=True
)

# 加载配置
config = load_config()
client = openai.OpenAI(api_key=config["openai_api_key"], base_url=config["openai_base_url"])
deepseek_api_key = config["deepseek_api_key"]
deepseek_endpoint = config["deepseek_api_url"]

# 初始化变量 (修复未定义问题)
topic = ""
chat_history = []
round_no = 0
style = "知乎式（分析理性）"  # 默认值
tone = "温和理性"  # 默认值
show_thoughts = False
show_chart = False
stance_mode = "参考立场 + 自主判断（推荐）"
chatgpt_stance = "无特定立场"
deepseek_stance = "无特定立场"

# 主内容区
with st.container():
    col1, col2 = st.columns([1, 1], gap="large")
    
    with col1:
        # 设置面板 - 苹果风格卡片
        st.markdown(
            """
            <div class="card">
                <div class="card-header">
                    <div class="icon">⚙️</div>
                    <h3>系统选项</h3>
                </div>
            """, 
            unsafe_allow_html=True
        )
        
        # 显示选项
        st.markdown("**基本选项**")
        show_thoughts = st.checkbox("显示思维过程（推理链）")
        show_chart = st.checkbox("显示评分图表")
        
        st.markdown('<div class="divider" style="margin: 20px 0;"></div>', unsafe_allow_html=True)
        
        # 立场设置
        st.markdown("**立场配置**")
        stance_mode = st.selectbox("立场引导模式：", [
            "参考立场 + 自主判断（推荐）",
            "强制立场（角色扮演）",
            "自由判断（无立场）"
        ])
        
        if stance_mode != "自由判断（无立场）":
            chatgpt_stance = st.text_input("ChatGPT 立场预设：", "无特定立场")
            deepseek_stance = st.text_input("Deepseek 立场预设：", "无特定立场")
        else:
            chatgpt_stance = "无特定立场"
            deepseek_stance = "无特定立场"
        
        st.markdown('</div>', unsafe_allow_html=True)  # 结束设置卡片
    
    with col2:
        # 议题管理 - 苹果风格卡片
        st.markdown(
            """
            <div class="card debate-card">  <!-- 添加特定class -->
                <div class="card-header">
                    <div class="icon">📚</div>
                    <h3>议题管理</h3>
                </div>
            """, 
            unsafe_allow_html=True
        )
        
        # 议题选择
        debates = list_saved_debates()
        debate_choice = st.selectbox("选择议题：", ["新建议题"] + debates, key="debate_select")
        
        # 议题管理按钮
        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            if st.button("🧹 清理无效议题", use_container_width=True, type="secondary"):
                cleaned = 0
                for name in debates:
                    try:
                        load_state(name)
                    except:
                        os.remove(f"data/{name}.json")
                        cleaned += 1
                st.success(f"已清理 {cleaned} 个异常议题文件")
        
        with col_btn2:
            if debate_choice != "新建议题":
                if st.button(f"🗑 删除当前议题", use_container_width=True, type="secondary"):
                    os.remove(f"data/{debate_choice}.json")
                    st.rerun()
        
        # 新建议题输入
        if debate_choice == "新建议题":
            topic = st.text_input("输入新议题：")
            if topic:
                if topic in debates:
                    topic, chat_history, round_no = load_state(topic)
                else:
                    round_no = 0
                    chat_history = []
                    chat_history.append({"role": "system", "content": f"你是 AI 辩论者 ChatGPT，{'请你基于知识与推理，自主判断' if stance_mode == '自由判断（无立场）' else f'请参考立场：{chatgpt_stance}'}。\n🗣 表达风格：{style}，🔥 激烈程度：{tone}，📌 议题：{topic}\n{('请输出你的思维链过程，逐步推理后再生成最终发言。' if show_thoughts else '')}"})
                    save_state(topic, chat_history, round_no)
        else:
            topic, chat_history, round_no = load_state(debate_choice)
        
        # 显示当前状态
        if "topic" in locals() or topic:
            st.markdown('<div class="divider" style="margin: 20px 0;"></div>', unsafe_allow_html=True)
            st.markdown(f"**当前议题:** {topic}")
            st.markdown(f"**已进行轮次:** {round_no}")
        
        st.markdown('</div>', unsafe_allow_html=True)  # 结束议题卡片

# 风格设置区域
st.markdown(
    """
    <div class="card style-card">  <!-- 添加特定class -->
        <div class="card-header">
            <div class="icon">🎭</div>
            <h3>辩论风格</h3>
        </div>
    """, 
    unsafe_allow_html=True
)

col_style1, col_style2 = st.columns(2)
with col_style1:
    tone = st.selectbox("激烈程度：", ["温和理性", "中度争锋", "激烈对立"])
with col_style2:
    style = st.selectbox("语言风格：", ["知乎式（分析理性）", "微博式（情绪化口号）", "Reddit式（讽刺+数据）", "Twitter式（短促有冲击）", "贴吧式（口水战/迷之逻辑）"])

st.markdown('</div>', unsafe_allow_html=True)  # 结束风格卡片

# 分隔线
st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

# 辩论控制区域
st.markdown(
    """
    <div class="card">
        <div class="card-header">
            <div class="icon">💬</div>
            <h3>辩论控制</h3>
        </div>
    """, 
    unsafe_allow_html=True
)

# 主辩论按钮容器
st.markdown('<div class="debate-button-container">', unsafe_allow_html=True)
debate_button = st.button("▶️ 开始辩论 / 继续下一轮", key="start_debate", 
                         help="点击开始新一轮辩论", use_container_width=True)
st.markdown('</div>', unsafe_allow_html=True)

st.markdown('</div>', unsafe_allow_html=True)  # 结束辩论控制卡片

# 定义AI回复函数 (修复变量未定义问题)
def chatgpt_reply(messages, style, tone, topic, stance_mode, chatgpt_stance, show_thoughts):
    system_message = f"你是 AI 辩论者 ChatGPT，{'请你基于知识与推理，自主判断' if stance_mode == '自由判断（无立场）' else f'请参考立场：{chatgpt_stance}'}。\n🗣 表达风格：{style}，🔥 激烈程度：{tone}，📌 议题：{topic}\n{('请输出你的思维链过程，逐步推理后再生成最终发言。' if show_thoughts else '')}"
    
    # 创建消息的副本，避免修改原始消息
    messages_copy = messages.copy()
    
    # 更新系统消息
    if len(messages_copy) > 0 and messages_copy[0]['role'] == 'system':
        messages_copy[0]['content'] = system_message
    else:
        messages_copy.insert(0, {'role': 'system', 'content': system_message})
    
    response = client.chat.completions.create(
        model=config["openai_model"],
        messages=messages_copy,
        temperature=0.7
    )
    return response.choices[0].message.content

def deepseek_reply(messages, style, tone, topic, stance_mode, deepseek_stance, show_thoughts):
    system_message = f"你是 AI 辩论者 Deepseek，{'请你基于知识与推理，自主判断' if stance_mode == '自由判断（无立场）' else f'请参考立场：{deepseek_stance}'}。\n🗣 表达风格：{style}，🔥 激烈程度：{tone}，📌 议题：{topic}\n{('请输出你的思维链过程，逐步推理后再生成最终发言。' if show_thoughts else '')}"
    
    # 创建消息的副本，避免修改原始消息
    messages_copy = messages.copy()
    
    # 更新系统消息
    if len(messages_copy) > 0 and messages_copy[0]['role'] == 'system':
        messages_copy[0]['content'] = system_message
    else:
        messages_copy.insert(0, {'role': 'system', 'content': system_message})
    
    headers = {
        "Authorization": f"Bearer {deepseek_api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "deepseek-chat",
        "messages": messages_copy,
        "temperature": 0.7
    }
    response = requests.post(deepseek_endpoint, headers=headers, json=payload)
    return response.json()['choices'][0]['message']['content']

# 处理辩论按钮点击
if debate_button:
    if not topic:
        st.warning("请先输入或选择议题")
        st.stop()
    
    with st.spinner(f"第 {round_no+1} 轮辩论生成中..."):
        # ChatGPT发言
        gpt_reply = chatgpt_reply(
            chat_history, 
            style, tone, topic, stance_mode, chatgpt_stance, show_thoughts
        )
        chat_history.append({"role": "assistant", "content": gpt_reply, "speaker": "ChatGPT"})

        # 为DeepSeek创建临时消息列表
        deepseek_messages = chat_history.copy()
        deepseek_messages.append({
            "role": "user",
            "content": f"请反驳上一位 AI 的观点。请以「{tone}」风格和「{style}」表达进行。"
        })
        
        # DeepSeek发言
        deepseek_resp = deepseek_reply(
            deepseek_messages, 
            style, tone, topic, stance_mode, deepseek_stance, show_thoughts
        )
        chat_history.append({"role": "assistant", "content": deepseek_resp, "speaker": "DeepSeek"})

        # 展示辩论内容
        st.markdown(
            f"""
            <div class="round-container">
                <div class="round-header">
                    <div>
                        <div class="round-title">第 {round_no+1} 轮辩论</div>
                        <div class="round-subtitle">风格: {style} | 激烈度: {tone}</div>
                    </div>
                </div>
                <div class="debate-content">
            """,
            unsafe_allow_html=True
        )
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown(f"#### 🤖 ChatGPT")
            st.write(gpt_reply)
            st.markdown('</div>', unsafe_allow_html=True)
            
        with col2:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown(f"#### 🧠 Deepseek")
            st.write(deepseek_resp)
            st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown('</div></div>', unsafe_allow_html=True)  # 结束辩论内容

        # 裁判点评与评分
        summary, score = get_critique_and_score(gpt_reply, deepseek_resp)
        
        # 使用Streamlit原生组件代替HTML字符串
        with st.container():
            st.subheader("裁判点评与评分")
            
            # 点评摘要卡片
            with st.expander("📝 点评摘要", expanded=True):
                st.markdown(f"<div style='padding: 12px;'>{summary}</div>", unsafe_allow_html=True)
            
            # 评分卡片布局
            col_score1, col_score2 = st.columns(2)
            
            with col_score1:
                st.markdown(
                    f"""
                    <div class="score-card chatgpt-score">
                        <div class="score-label">ChatGPT 得分</div>
                        <div class="score-value">{score['ChatGPT']}</div>
                        <div class="score-label">/10</div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            
            with col_score2:
                st.markdown(
                    f"""
                    <div class="score-card deepseek-score">
                        <div class="score-label">Deepseek 得分</div>
                        <div class="score-value">{score['Deepseek']}</div>
                        <div class="score-label">/10</div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

        # 评分图表
        if show_chart:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown("#### 📊 评分对比")
            fig, ax = plt.subplots(figsize=(8, 4))
            
            # 苹果风格的图表
            colors = ['#2c6fbb', '#5a4fcf']
            bars = ax.bar(['ChatGPT', 'Deepseek'], [score['ChatGPT'], score['Deepseek']], color=colors)
            ax.set_ylim(0, 10)
            ax.set_ylabel('分数', fontsize=12)
            ax.grid(axis='y', linestyle='--', alpha=0.7)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            
            # 在柱子上方添加数值
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                        f'{height}', ha='center', va='bottom', fontsize=12)
            
            st.pyplot(fig)
            st.markdown('</div>', unsafe_allow_html=True)

        save_state(topic, chat_history, round_no + 1)

    # 准备导出文档
    doc = Document()
    doc.add_heading(f"AI Debate - {topic}", 0)
    doc.add_paragraph(f"议题：{topic}")
    doc.add_paragraph(f"风格: {style} | 激烈度: {tone}")
    doc.add_heading("辩论过程：", level=1)
    
    # 只提取辩论内容
    debate_content = [msg for msg in chat_history if msg["role"] == "assistant"]
    total_rounds = len(debate_content) // 2
    
    for i in range(total_rounds):
        gpt_round = debate_content[i * 2].get("content", "(缺失)")
        deepseek_round = debate_content[i * 2 + 1].get("content", "(缺失)")
        doc.add_paragraph(f"[ChatGPT 第{i+1}轮]：{gpt_round}")
        doc.add_paragraph(f"[Deepseek 第{i+1}轮]：{deepseek_round}")
        summary, score = get_critique_and_score(gpt_round, deepseek_round)
        doc.add_paragraph(f"[点评员总结]：{summary}（评分：ChatGPT {score['ChatGPT']}/10, Deepseek {score['Deepseek']}/10）")

    output_path = f"output/Debate_{topic[:20].replace(' ', '_')}.docx"
    doc.save(output_path)
    
    # 下载按钮
    st.markdown('<div class="card">', unsafe_allow_html=True)
    with open(output_path, "rb") as f:
        st.download_button("📄 下载完整辩论文档", f, file_name=os.path.basename(output_path),
                          use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

# 完整辩论记录区域
st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

# 可折叠的完整辩论记录
with st.expander("📜 查看完整辩论记录", expanded=False):
    st.markdown(
        """
        <div class="card">
            <div class="card-header">
                <div class="icon">🗂</div>
                <h3>完整辩论记录</h3>
            </div>
        """, 
        unsafe_allow_html=True
    )
    
    if len(chat_history) > 1:
        # 只显示辩论内容
        debate_content = [msg for msg in chat_history if msg["role"] == "assistant"]
        
        if len(debate_content) == 0:
            st.markdown(
                """
                <div class="empty-state">
                    <div class="empty-state-icon">📭</div>
                    <p>暂无辩论记录</p>
                </div>
                """, 
                unsafe_allow_html=True
            )
        else:
            for i in range(0, len(debate_content), 2):
                round_num = (i // 2) + 1
                
                st.markdown(
                    f"""
                    <div class="round-container">
                        <div class="round-header">
                            <div class="round-title">第 {round_num} 轮</div>
                        </div>
                        <div class="debate-content">
                    """,
                    unsafe_allow_html=True
                )
                
                col_rec1, col_rec2 = st.columns(2)
                with col_rec1:
                    st.markdown('<div class="card">', unsafe_allow_html=True)
                    st.markdown(f"#### 🤖 ChatGPT")
                    st.write(debate_content[i].get("content", "(缺失)"))
                    st.markdown('</div>', unsafe_allow_html=True)
                
                with col_rec2:
                    if i + 1 < len(debate_content):
                        st.markdown('<div class="card">', unsafe_allow_html=True)
                        st.markdown(f"#### 🧠 Deepseek")
                        st.write(debate_content[i+1].get("content", "(缺失)"))
                        st.markdown('</div>', unsafe_allow_html=True)
                
                st.markdown('</div></div>', unsafe_allow_html=True)  # 结束辩论内容
    else:
        st.markdown(
            """
            <div class="empty-state">
                <div class="empty-state-icon">📭</div>
                <p>暂无辩论记录</p>
            </div>
            """, 
            unsafe_allow_html=True
        )
    
    st.markdown('</div>', unsafe_allow_html=True)  # 结束完整记录卡片

# 页脚
st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
st.markdown(
    """
    <div class="footer">
        <p>© 2023 AI 辩论系统 | 基于 ChatGPT 和 DeepSeek AI 技术 | v1.6.0</p>
        <p>苹果设计风格 | 简洁 · 优雅 · 强大</p>
    </div>
    """, 
    unsafe_allow_html=True
)